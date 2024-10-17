import docx
from datetime import date

mydoc = docx.Document()
daysInMonths = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
seniors = ["Madhumithaa", "Bhavana", "Jayalaxmi", "Marjeena", "Rashmi", "Shruthi", "Srushti", "Adyashree", "Nidhish", "Shilpy", "Vidyashree"]
juniors = ["Jharna", "Malini", "Pratiksha", "Shweta", "Aishwarya", "Harshita"]

current_month = 10
senior_start_index = 8
junior_start_index = 4
rows = daysInMonths[current_month - 1] + 1
columns = 2
night_duty_table = mydoc.add_table(rows, columns)

#night_duty_table.add_heading("NIGHT DUTY")
first_row = night_duty_table.rows[0]
first_row.cells[0].text = 'DATE'
first_row.cells[1].text = 'PG SCHOLAR NIGHT DUTY'

for x in range(rows - 1):
	dutyDate = date(2024, current_month, x + 1)
	row = night_duty_table.rows[x + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print(dutyDate.strftime('%d-%m-%Y'))
	seniorDoctor = seniors[(x + senior_start_index) % 11]
	juniorDoctor = juniors[(x + junior_start_index) % 6]
	row.cells[1].text = 'Dr. ' + seniorDoctor + ', Dr. ' + juniorDoctor

mydoc.save('C:\\Personal\\Python\\output\\dutychart.docx')

