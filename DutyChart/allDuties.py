import docx
from datetime import date, timedelta, datetime
import json
from pathlib import Path
from docx.shared import Cm

with open("config.json") as config_file:
	configurations = json.load(config_file)

daysInMonths = configurations["daysInMonths"]
holidays = configurations["Holidays"]
thirdYears = configurations["ThirdYears"]
secondYears = configurations["SecondYears"]
firstYears = configurations["FirstYears"]
babyJuniors = configurations["BabyJuniors"]

current_month = int(configurations["current_month"])
thirdYearsIndexNight = int(configurations["ThirdYearsStartPositionNight"]) - 1
secondYearsIndexNight = int(configurations["SecondYearsStartPositionNight"]) - 1
firstYearsIndexNight = int(configurations["FirstYearsStartPositionNight"]) - 1
babyJuniorsIndexNight = int(configurations["BabyJuniorsStartPositionNight"]) - 1

thirdYearsIndexHoliday = int(configurations["ThirdYearsStartPositionHoliday"]) - 1
secondYearsIndexHoliday = int(configurations["SecondYearsStartPositionHoliday"]) - 1
firstYearsIndexHoliday = int(configurations["FirstYearsStartPositionHoliday"]) - 1
babyJuniorsIndexHoliday = int(configurations["BabyJuniorsStartPositionHoliday"]) - 1

thirdYearsIndexEve = int(configurations["ThirdYearsStartPositionEve"]) - 1
secondYearsIndexEve = int(configurations["SecondYearsStartPositionEve"]) - 1
firstYearsIndexEve = int(configurations["FirstYearsStartPositionEve"]) - 1
babyJuniorsIndexEve = int(configurations["BabyJuniorsStartPositionEve"]) - 1

def isAvailable(startDate, endDate, todaysDate):
	if (todaysDate >= startDate and todaysDate <= endDate):
		return False
	return True

def isNotAvailable(startDate, endDate, todaysDate):
	return not(isAvailable(startDate, endDate, todaysDate))

rows = daysInMonths[current_month - 1] + 1
columns = 2

mydoc = docx.Document()
night_duty_table = mydoc.add_table(rows, columns)
night_duty_table.style = 'Table Grid'

#night_duty_table.add_heading("NIGHT DUTY")
first_row = night_duty_table.rows[0]
first_row.cells[0].text = 'DATE'
first_row.cells[1].text = 'PG SCHOLAR NIGHT DUTY'
first_row.cells[0].width = Cm(2.5)
first_row.cells[1].width = Cm(5)

addComma = False
doThirdYears = False
doSecondYears = False
if (configurations["StartWithSecondYear"].lower() == "yes"):
	doSecondYears = True


for x in range(rows - 1):
	dutyDate = date(2024, current_month, x + 1)
	row = night_duty_table.rows[x + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print(dutyDate.strftime('%d-%m-%Y'))
	doctor=""
	if ( (configurations["BothThirdYearsAndSecondYearsPresent"]).lower() == "no"):
		if (configurations["ThirdYearAbscentEndDate"] < dutyDate.day):
			doSecondYears = False
			doThirdYears = True
		if (configurations["ThirdYearAbscentEndDate"] >= dutyDate.day):
			doSecondYears = True
			doThirdYears = False

		if (doSecondYears):
			if (isAvailable(int(configurations["SecondYearAbscentStartDate"]), int(configurations["SecondYearAbscentEndDate"]), dutyDate.day)):
				doctor += ("Dr. " + secondYears[secondYearsIndexNight])
				secondYearsIndexNight += 1
				addComma = True
				if (secondYearsIndexNight == len(secondYears)):
					doSecondYears = False
					doThirdYears = True
		elif (doThirdYears):
			if (isAvailable(int(configurations["ThirdYearAbscentStartDate"]), int(configurations["ThirdYearAbscentEndDate"]), dutyDate.day)):
				doctor += ("Dr. " + thirdYears[thirdYearsIndexNight])
				thirdYearsIndexNight += 1
				addComma = True
				if (thirdYearsIndexNight == len(thirdYears)):
					doThirdYears = False
					doSecondYears = True
	else:
		if (isAvailable(int(configurations["ThirdYearAbscentStartDate"]), int(configurations["ThirdYearAbscentEndDate"]), dutyDate.day)):
			doctor += ("Dr. " + thirdYears[thirdYearsIndexNight])
			thirdYearsIndexNight += 1
			addComma = True

		if (isAvailable(int(configurations["SecondYearAbscentStartDate"]), int(configurations["SecondYearAbscentEndDate"]), dutyDate.day)):
			doctor += ("Dr. " + secondYears[secondYearsIndexNight])
			secondYearsIndexNight += 1
			addComma = True

	if (isAvailable(int(configurations["FirstYearAbscentStartDate"]), int(configurations["FirstYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + firstYears[firstYearsIndexNight])
		firstYearsIndexNight += 1

	if (isAvailable(int(configurations["BabyJuniorsAbscentStartDate"]), int(configurations["BabyJuniorsAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + babyJuniors[babyJuniorsIndexNight])
		babyJuniorsIndexNight += 1
		babyJuniorsIndexNight %= len(babyJuniors)

	row.cells[1].text = doctor
	row.cells[0].width = Cm(2.5)
	row.cells[1].width = Cm(3)
	thirdYearsIndexNight %= len(thirdYears)
	secondYearsIndexNight %= len(secondYears)
	firstYearsIndexNight %= len(firstYears)

mydoc.add_paragraph("\n")

# afternoon
afternoonDutyStartDate = datetime.strptime(configurations["AfternoonDutyStartDate"], "%d/%m/%Y").date()
afternoonDutyEndDate = afternoonDutyStartDate + timedelta(days=6)
noOfDays = daysInMonths[current_month - 1]
if (afternoonDutyStartDate.month != current_month):
	prevMonthLen = daysInMonths[(current_month - 1 ) % 12]
	noOfDays += prevMonthLen - afternoonDutyStartDate.day + 1

afternoonDutyStartIndex = configurations["AfternoonDutyStartPosition"] - 1
rows = noOfDays//7 + 2
columns = 3
afternoon_duty_table = mydoc.add_table(rows, columns)
afternoon_duty_table.style = 'Table Grid'
first_row = afternoon_duty_table.rows[0]
first_row.cells[0].text = 'AFTERNOON DUTY'
first_row.cells[1].text = '12:30-13:00'
first_row.cells[2].text = '13:00-13:30'
first_row.cells[0].width = Cm(3.5)
first_row.cells[1].width = Cm(3.5)
first_row.cells[2].width = Cm(3.5)

afternoonDutyDoctors=firstYears
if (configurations["AfternoonDutyDoneByWhichYear"] == 2):
	afternoonDutyDoctors=secondYears


for index in range(rows - 1):
	row = afternoon_duty_table.rows[index + 1]
	row.cells[0].text = afternoonDutyStartDate.strftime("%d-%m-%Y") + " - " + afternoonDutyEndDate.strftime("%d-%m-%Y")
	# print(dutyDate.strftime('%d-%m-%Y'))
	doctors=""
	row.cells[1].text = 'Dr. ' + afternoonDutyDoctors[afternoonDutyStartIndex]
	afternoonDutyStartIndex += 1
	afternoonDutyStartIndex %= len(afternoonDutyDoctors)

	row.cells[2].text = 'Dr. ' + afternoonDutyDoctors[afternoonDutyStartIndex]
	afternoonDutyStartIndex += 1
	afternoonDutyStartIndex %= len(afternoonDutyDoctors)
	row.cells[0].width = Cm(3.5)
	row.cells[1].width = Cm(3.5)
	row.cells[2].width = Cm(3.5)
	afternoonDutyStartDate = afternoonDutyStartDate + timedelta(days=7)
	afternoonDutyEndDate = afternoonDutyEndDate + timedelta(days=7)

# evening
mydoc.add_paragraph("\n")
rows = daysInMonths[current_month - 1] - len(holidays) + 1
columns = 2
index = 0
date_of_month = 0
evening_duty_table = mydoc.add_table(rows, columns)
evening_duty_table.style = 'Table Grid'
first_row = evening_duty_table.rows[0]
first_row.cells[0].text = 'DATE'
first_row.cells[1].text = 'PG SCHOLAR EVENING DUTY'
first_row.cells[0].width = Cm(2.5)
first_row.cells[1].width = Cm(5.5)
count = 1
while index < rows - 1:
	date_of_month += 1
	if (holidays.count(date_of_month) == 1):
		continue
	dutyDate = date(2024, current_month, date_of_month)
	row = evening_duty_table.rows[index + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print("date of month = %s", date_of_month)
	doctor=""
	addComma = False
	if (isAvailable(int(configurations["ThirdYearAbscentStartDate"]), int(configurations["ThirdYearAbscentEndDate"]), dutyDate.day)):
		doctor += ("Dr. " + thirdYears[thirdYearsIndexEve])
		if (count == 2):
			thirdYearsIndexEve += 1
		addComma = True

	if (isAvailable(int(configurations["SecondYearAbscentStartDate"]), int(configurations["SecondYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + secondYears[secondYearsIndexEve])
		if (count == 2):
			secondYearsIndexEve += 1
		addComma = True

	if (isAvailable(int(configurations["FirstYearAbscentStartDate"]), int(configurations["FirstYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		if (count == 2):
			doctor += ("Dr. " + firstYears[firstYearsIndexEve])
		firstYearsIndexEve += 1

	if (isAvailable(int(configurations["BabyJuniorsAbscentStartDate"]), int(configurations["BabyJuniorsAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + babyJuniors[babyJuniorsIndexEve])
		if (count == 2):
			babyJuniorsIndexEve += 1
		babyJuniorsIndexEve %= len(babyJuniors)

	row.cells[1].text = doctor
	thirdYearsIndexEve %= len(thirdYears)
	secondYearsIndexEve %= len(secondYears)
	firstYearsIndexEve %= len(firstYears)
	index += 1
	count += 1
	if (count == 3):
		count = 1
	row.cells[0].width = Cm(2.5)
	row.cells[1].width = Cm(5.5)


#holiday_duty_table.add_heading("HOLIDAY DUTY")
mydoc.add_paragraph("\n")
rows = len(holidays) + 1
columns = 3
index = 0
holiday_duty_table = mydoc.add_table(rows, columns)
holiday_duty_table.style = 'Table Grid'

first_row = holiday_duty_table.rows[0]
first_row.cells[0].text = 'HOLIDAY DUTY'
first_row.cells[1].text = 'MORNING SHIFT'
first_row.cells[2].text = 'AFTERNOON SHIFT'
first_row.cells[0].width = Cm(7)
first_row.cells[1].width = Cm(7)
first_row.cells[2].width = Cm(7)

index = 0

for x in holidays:
	dutyDate = date(2024, current_month, x)
	row = holiday_duty_table.rows[index + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print(dutyDate.strftime('%d-%m-%Y'))
	doctor = ""
	addComma = False
	if (isAvailable(int(configurations["ThirdYearAbscentStartDate"]), int(configurations["ThirdYearAbscentEndDate"]), dutyDate.day)):
		doctor += ("Dr. " + thirdYears[thirdYearsIndexHoliday])
		thirdYearsIndexHoliday += 1
		addComma = True

	if (isAvailable(int(configurations["SecondYearAbscentStartDate"]), int(configurations["SecondYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + secondYears[secondYearsIndexHoliday])
		secondYearsIndexHoliday += 1
		addComma = True

	if (isAvailable(int(configurations["FirstYearAbscentStartDate"]), int(configurations["FirstYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + firstYears[firstYearsIndexHoliday])
		firstYearsIndexHoliday += 1

	if (isAvailable(int(configurations["BabyJuniorsAbscentStartDate"]), int(configurations["BabyJuniorsAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + babyJuniors[babyJuniorsIndexHoliday])
		babyJuniorsIndexHoliday += 1
		babyJuniorsIndexHoliday %= len(babyJuniors)

	row.cells[1].text = doctor
	thirdYearsIndexHoliday %= len(thirdYears)
	secondYearsIndexHoliday %= len(secondYears)
	firstYearsIndexHoliday %= len(firstYears)

	doctor = ""
	addComma = False
	if (isAvailable(int(configurations["ThirdYearAbscentStartDate"]), int(configurations["ThirdYearAbscentEndDate"]), dutyDate.day)):
		doctor += ("Dr. " + thirdYears[thirdYearsIndexHoliday])
		thirdYearsIndexHoliday += 1
		addComma = True

	if (isAvailable(int(configurations["SecondYearAbscentStartDate"]), int(configurations["SecondYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + secondYears[secondYearsIndexHoliday])
		secondYearsIndexHoliday += 1
		addComma = True

	if (isAvailable(int(configurations["FirstYearAbscentStartDate"]), int(configurations["FirstYearAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + firstYears[firstYearsIndexHoliday])
		firstYearsIndexHoliday += 1

	if (isAvailable(int(configurations["BabyJuniorsAbscentStartDate"]), int(configurations["BabyJuniorsAbscentEndDate"]), dutyDate.day)):
		if (addComma):
			doctor += ", "
		doctor += ("Dr. " + babyJuniors[babyJuniorsIndexHoliday])
		babyJuniorsIndexHoliday += 1
		babyJuniorsIndexHoliday %= len(babyJuniors)

	row.cells[2].text = doctor
	thirdYearsIndexHoliday %= len(thirdYears)
	secondYearsIndexHoliday %= len(secondYears)
	firstYearsIndexHoliday %= len(firstYears)

	index += 1
	row.cells[0].width = Cm(7)
	row.cells[1].width = Cm(7)
	row.cells[2].width = Cm(7)


Path("output").mkdir(parents=False, exist_ok=True)
mydoc.save('output\\dutychart.docx')

