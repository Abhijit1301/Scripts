import docx
from datetime import date

mydoc = docx.Document()
daysInMonths = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
seniors = ["Madhumithaa", "Bhavana", "Jayalaxmi", "Marjeena", "Rashmi", "Shruthi", "Srushti", "Adyashree", "Nidhish", "Shilpy", "Vidyashree"]
juniors = ["Jharna", "Malini", "Pratiksha", "Shweta", "Aishwarya", "Harshita"]
holidays = [2, 6, 11, 12, 13, 17, 20, 27, 31]

current_month = 10
senior_start_index = 2
junior_start_index = 2
rows = len(holidays) + 1
columns = 3
holiday_duty_table = mydoc.add_table(rows, columns)

#holiday_duty_table.add_heading("HOLIDAY DUTY")
first_row = holiday_duty_table.rows[0]
first_row.cells[0].text = 'DATE'
first_row.cells[1].text = 'MORNING SHIFT'
first_row.cells[2].text = 'AFTERNOON SHIFT'

index = 0

for x in holidays:
	dutyDate = date(2024, current_month, x)
	row = holiday_duty_table.rows[index + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print(dutyDate.strftime('%d-%m-%Y'))
	seniorDoctor = seniors[senior_start_index % 11]
	juniorDoctor = juniors[junior_start_index % 6]
	row.cells[1].text = 'Dr. ' + seniorDoctor + ', Dr. ' + juniorDoctor

	senior_start_index += 1
	junior_start_index += 1
	seniorDoctor = seniors[senior_start_index % 11]
	juniorDoctor = juniors[junior_start_index % 6]
	row.cells[2].text = 'Dr. ' + seniorDoctor + ', Dr. ' + juniorDoctor

	senior_start_index += 1
	junior_start_index += 1
	index += 1

mydoc.save('C:\\Personal\\Python\\output\\holiday_dutychart.docx')

