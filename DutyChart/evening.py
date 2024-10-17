import docx
from datetime import date

mydoc = docx.Document()
daysInMonths = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
seniors = ["Madhumithaa", "Bhavana", "Jayalaxmi", "Marjeena", "Rashmi", "Shruthi", "Srushti", "Adyashree", "Nidhish", "Shilpy", "Vidyashree"]
juniors = ["Jharna", "Malini", "Pratiksha", "Shweta", "Aishwarya", "Harshita"]
holidays = [2, 6, 11, 12, 13, 17, 20, 27, 31]

current_month = 10
senior_index = 9
junior_index = 0
senior_reps = 2
junior_reps = 1
rows = daysInMonths[current_month - 1] - len(holidays) + 1
columns = 2
evening_duty_table = mydoc.add_table(rows, columns)

#evening_duty_table.add_heading("NIGHT DUTY")
first_row = evening_duty_table.rows[0]
first_row.cells[0].text = 'DATE'
first_row.cells[1].text = 'PG SCHOLAR NIGHT DUTY'
date_of_month = 0
index = 0
while index < rows - 1:
	date_of_month += 1
	if (holidays.count(date_of_month) == 1):
		continue
	# print("date of month = %s", date_of_month)
	dutyDate = date(2024, current_month, date_of_month)
	row = evening_duty_table.rows[index + 1]
	row.cells[0].text = dutyDate.strftime('%d-%m-%Y')
	# print(dutyDate.strftime('%d-%m-%Y'))
	seniorDoctor = seniors[senior_index % 11]
	juniorDoctor = juniors[junior_index % 6]
	row.cells[1].text = 'Dr. ' + seniorDoctor + ', Dr. ' + juniorDoctor
	index += 1
	senior_reps -= 1
	junior_reps -= 1
	if (senior_reps == 0):
		senior_reps = 2
		senior_index += 1
	if (junior_reps == 0):
		junior_reps = 2
		junior_index +=1

mydoc.save('C:\\Personal\\Python\\output\\evening_dutychart.docx')

